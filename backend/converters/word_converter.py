"""
Word Converter Module
Handles conversions from DOCX to other formats
"""

import os
import re
from pathlib import Path
from backend.converters.base_converter import BaseConverter, ConversionError

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    import docx2pdf
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False

try:
    import subprocess as _sp
    _kw = {}
    if os.name == 'nt':
        _kw['creationflags'] = 0x08000000  # CREATE_NO_WINDOW
    _sp.run(['pandoc', '--version'], capture_output=True, timeout=5, **_kw)
    PANDOC_AVAILABLE = True
except Exception:
    PANDOC_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False


class WordConverter(BaseConverter):
    """Convert Word (DOCX) files to various formats"""

    SUPPORTED_OUTPUTS = ['pdf', 'txt', 'md', 'html']

    def convert(self, input_path: str, output_path: str, target_format: str) -> bool:
        self.check_cancel()
        if target_format == 'pdf':
            return self._to_pdf(input_path, output_path)
        elif target_format in ['txt', 'md']:
            return self._to_text(input_path, output_path, target_format)
        elif target_format == 'html':
            return self._to_html(input_path, output_path)
        else:
            raise ConversionError(f"Unsupported target format: {target_format}")

    def _to_pdf(self, input_path: str, output_path: str) -> bool:
        self.report_progress(10, "Converting Word to PDF...")

        # Method 1: Direct Word COM automation (best quality, preserves all styling)
        if os.name == 'nt':
            try:
                import win32com.client
                self.report_progress(15, "Opening Microsoft Word...")
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                word.DisplayAlerts = False
                abs_in = os.path.abspath(input_path)
                abs_out = os.path.abspath(output_path)
                doc = word.Documents.Open(abs_in, ReadOnly=True)
                self.report_progress(50, "Saving as PDF...")
                doc.SaveAs2(abs_out, FileFormat=17)  # 17 = wdFormatPDF
                doc.Close(False)
                word.Quit()
                if os.path.exists(output_path):
                    self.report_progress(100, "Done")
                    return True
            except Exception:
                pass

        # Method 2: Pandoc
        if PANDOC_AVAILABLE:
            try:
                self.report_progress(30, "Converting via Pandoc...")
                import subprocess
                kw = {}
                if os.name == 'nt':
                    kw['creationflags'] = 0x08000000
                result = subprocess.run(
                    ['pandoc', input_path, '-o', output_path],
                    capture_output=True, text=True, timeout=120, **kw
                )
                if result.returncode == 0 and os.path.exists(output_path):
                    self.report_progress(100, "Done")
                    return True
            except Exception:
                pass

        # Method 3: LibreOffice
        try:
            self.report_progress(50, "Trying LibreOffice...")
            return self._convert_with_libreoffice(input_path, output_path)
        except Exception:
            pass

        # Method 4: Pure Python (reportlab + python-docx)
        if REPORTLAB_AVAILABLE and PYTHON_DOCX_AVAILABLE:
            self.report_progress(60, "Converting via reportlab...")
            return self._convert_with_reportlab(input_path, output_path)

        raise ConversionError(
            "DOCX to PDF failed — no conversion method available",
            "Install Pandoc for best results"
        )

    def _convert_with_libreoffice(self, input_path: str, output_path: str) -> bool:
        import subprocess
        output_dir = os.path.dirname(output_path)
        cmd = [
            'libreoffice', '--headless', '--convert-to', 'pdf',
            '--outdir', output_dir, input_path
        ]
        kw = {}
        if os.name == 'nt':
            kw['creationflags'] = 0x08000000
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60, **kw)

        if result.returncode == 0:
            base_name = Path(input_path).stem
            expected_output = os.path.join(output_dir, f"{base_name}.pdf")
            if os.path.exists(expected_output) and expected_output != output_path:
                os.rename(expected_output, output_path)
            self.report_progress(100, "Done")
            return True
        raise ConversionError("LibreOffice conversion failed")

    def _convert_with_reportlab(self, input_path: str, output_path: str) -> bool:
        doc = Document(input_path)
        pdf = SimpleDocTemplate(output_path, pagesize=letter,
                                leftMargin=72, rightMargin=72,
                                topMargin=72, bottomMargin=72)
        styles = getSampleStyleSheet()
        story = []

        for para in doc.paragraphs:
            self.check_cancel()
            text = para.text.strip()
            if not text:
                story.append(Spacer(1, 6))
                continue

            safe_text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            style_name = para.style.name.lower() if para.style else ""

            if 'heading 1' in style_name:
                story.append(Paragraph(safe_text, styles['Heading1']))
            elif 'heading 2' in style_name:
                story.append(Paragraph(safe_text, styles['Heading2']))
            elif 'heading 3' in style_name:
                story.append(Paragraph(safe_text, styles['Heading3']))
            else:
                story.append(Paragraph(safe_text, styles['Normal']))

        for table in doc.tables:
            self.check_cancel()
            data = []
            for row in table.rows:
                data.append([cell.text.strip() for cell in row.cells])
            if data:
                t = Table(data)
                t.setStyle(TableStyle([
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e0e0e0')),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('PADDING', (0, 0), (-1, -1), 4),
                ]))
                story.append(Spacer(1, 12))
                story.append(t)

        if not story:
            story.append(Paragraph("(Empty document)", styles['Normal']))

        pdf.build(story)
        self.report_progress(100, "Done")
        return True

    def _extract_doc_content(self, input_path):
        """Extract paragraphs and tables from DOCX."""
        if not PYTHON_DOCX_AVAILABLE:
            raise ConversionError("python-docx not installed", "pip install python-docx")
        return Document(input_path)

    def _to_text(self, input_path: str, output_path: str, format_type: str) -> bool:
        try:
            self.report_progress(10, "Reading Word document...")
            doc = self._extract_doc_content(input_path)
            content_parts = []

            total = len(doc.paragraphs) + len(doc.tables)
            done = 0

            for para in doc.paragraphs:
                self.check_cancel()
                text = para.text.strip()
                if not text:
                    continue

                if format_type == 'md':
                    style_name = para.style.name.lower() if para.style else ""
                    if 'heading 1' in style_name:
                        content_parts.append(f"# {text}")
                    elif 'heading 2' in style_name:
                        content_parts.append(f"## {text}")
                    elif 'heading 3' in style_name:
                        content_parts.append(f"### {text}")
                    elif 'heading' in style_name:
                        level = re.search(r'\d+', style_name)
                        hashes = '#' * min(int(level.group()), 6) if level else '##'
                        content_parts.append(f"{hashes} {text}")
                    else:
                        formatted = ""
                        for run in para.runs:
                            rt = run.text
                            if run.bold and run.italic:
                                rt = f"***{rt}***"
                            elif run.bold:
                                rt = f"**{rt}**"
                            elif run.italic:
                                rt = f"*{rt}*"
                            formatted += rt
                        content_parts.append(formatted if formatted else text)
                else:
                    content_parts.append(text)

                done += 1
                self.report_progress(10 + int(80 * done / max(total, 1)))

            for table in doc.tables:
                self.check_cancel()
                if format_type == 'md':
                    content_parts.append("")
                    for i, row in enumerate(table.rows):
                        row_text = "| " + " | ".join(
                            cell.text.strip().replace('\n', ' ') for cell in row.cells
                        ) + " |"
                        content_parts.append(row_text)
                        if i == 0:
                            sep = "|" + "|".join(" --- " for _ in row.cells) + "|"
                            content_parts.append(sep)
                    content_parts.append("")
                else:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        content_parts.append(row_text)
                    content_parts.append("-" * 50)
                done += 1

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write("\n\n".join(content_parts))
            self.report_progress(100, "Done")
            return True
        except ConversionError:
            raise
        except Exception as e:
            raise ConversionError(f"DOCX to text failed: {e}")

    def _to_html(self, input_path: str, output_path: str) -> bool:
        try:
            self.report_progress(10, "Reading Word document...")
            doc = self._extract_doc_content(input_path)
            html_parts = []

            for para in doc.paragraphs:
                self.check_cancel()
                text = para.text.strip()
                if not text:
                    continue
                escaped = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                style_name = para.style.name.lower() if para.style else ""
                if 'heading 1' in style_name:
                    html_parts.append(f"<h1>{escaped}</h1>")
                elif 'heading 2' in style_name:
                    html_parts.append(f"<h2>{escaped}</h2>")
                elif 'heading 3' in style_name:
                    html_parts.append(f"<h3>{escaped}</h3>")
                else:
                    html_parts.append(f"<p>{escaped}</p>")

            self.report_progress(70, "Building HTML...")
            for table in doc.tables:
                self.check_cancel()
                html_parts.append("<table border='1' cellpadding='5'>")
                for i, row in enumerate(table.rows):
                    tag = "th" if i == 0 else "td"
                    cells = "".join(f"<{tag}>{c.text.strip()}</{tag}>" for c in row.cells)
                    html_parts.append(f"<tr>{cells}</tr>")
                html_parts.append("</table>")

            html = f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8"><title>Converted from Word</title>
<style>body{{font-family:Arial,sans-serif;max-width:800px;margin:40px auto;line-height:1.6;}}
table{{border-collapse:collapse;width:100%;margin:1em 0;}}</style>
</head><body>
{"".join(html_parts)}
</body></html>"""

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html)
            self.report_progress(100, "Done")
            return True
        except ConversionError:
            raise
        except Exception as e:
            raise ConversionError(f"DOCX to HTML failed: {e}")
