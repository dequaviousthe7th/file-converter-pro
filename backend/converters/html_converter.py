"""
HTML Converter Module
Handles conversions from HTML to other formats
"""

import re
from pathlib import Path
from backend.converters.base_converter import BaseConverter, ConversionError

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

try:
    from weasyprint import HTML as WeasyprintHTML
    WEASYPRINT_AVAILABLE = True
except Exception:
    WEASYPRINT_AVAILABLE = False

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False


class HTMLConverter(BaseConverter):
    """Convert HTML files to various formats"""

    SUPPORTED_OUTPUTS = ['pdf', 'docx', 'txt', 'md']

    def convert(self, input_path: str, output_path: str, target_format: str) -> bool:
        self.check_cancel()
        if target_format == 'pdf':
            return self._to_pdf(input_path, output_path)
        elif target_format == 'docx':
            return self._to_docx(input_path, output_path)
        elif target_format == 'txt':
            return self._to_text(input_path, output_path)
        elif target_format == 'md':
            return self._to_markdown(input_path, output_path)
        else:
            raise ConversionError(f"Unsupported target format: {target_format}")

    def _read_html(self, path):
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def _to_pdf(self, input_path: str, output_path: str) -> bool:
        self.report_progress(10, "Converting HTML to PDF...")
        if WEASYPRINT_AVAILABLE:
            try:
                self.check_cancel()
                WeasyprintHTML(filename=input_path).write_pdf(output_path)
                self.report_progress(100, "Done")
                return True
            except Exception:
                pass

        if REPORTLAB_AVAILABLE:
            return self._to_pdf_reportlab(input_path, output_path)

        raise ConversionError("No PDF library available for HTML conversion")

    def _to_pdf_reportlab(self, input_path: str, output_path: str) -> bool:
        text = self._extract_text(input_path)
        self.check_cancel()
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        for para in text.split('\n\n'):
            if para.strip():
                safe = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(safe.replace('\n', '<br/>'), styles['Normal']))
                story.append(Spacer(1, 12))
        doc.build(story)
        self.report_progress(100, "Done")
        return True

    def _to_docx(self, input_path: str, output_path: str) -> bool:
        if not PYTHON_DOCX_AVAILABLE:
            raise ConversionError("python-docx not installed")
        self.report_progress(10, "Converting HTML to Word...")
        html = self._read_html(input_path)
        self.check_cancel()

        if BS4_AVAILABLE:
            soup = BeautifulSoup(html, 'html.parser')
            doc = Document()
            for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'li']):
                self.check_cancel()
                text = tag.get_text(strip=True)
                if not text:
                    continue
                if tag.name == 'h1':
                    doc.add_heading(text, level=1)
                elif tag.name == 'h2':
                    doc.add_heading(text, level=2)
                elif tag.name in ('h3', 'h4'):
                    doc.add_heading(text, level=3)
                else:
                    doc.add_paragraph(text)
            doc.save(output_path)
        else:
            text = self._strip_tags(html)
            doc = Document()
            for para in text.split('\n\n'):
                if para.strip():
                    doc.add_paragraph(para.strip())
            doc.save(output_path)

        self.report_progress(100, "Done")
        return True

    def _to_text(self, input_path: str, output_path: str) -> bool:
        self.report_progress(10, "Extracting text from HTML...")
        text = self._extract_text(input_path)
        self.check_cancel()
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
        self.report_progress(100, "Done")
        return True

    def _to_markdown(self, input_path: str, output_path: str) -> bool:
        self.report_progress(10, "Converting HTML to Markdown...")
        html = self._read_html(input_path)
        self.check_cancel()

        if BS4_AVAILABLE:
            soup = BeautifulSoup(html, 'html.parser')
            md_parts = []
            for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'li', 'pre', 'code', 'a']):
                text = tag.get_text(strip=True)
                if not text:
                    continue
                if tag.name == 'h1':
                    md_parts.append(f"# {text}")
                elif tag.name == 'h2':
                    md_parts.append(f"## {text}")
                elif tag.name == 'h3':
                    md_parts.append(f"### {text}")
                elif tag.name == 'h4':
                    md_parts.append(f"#### {text}")
                elif tag.name == 'li':
                    md_parts.append(f"- {text}")
                elif tag.name in ('pre', 'code'):
                    md_parts.append(f"```\n{text}\n```")
                elif tag.name == 'a':
                    href = tag.get('href', '')
                    md_parts.append(f"[{text}]({href})")
                else:
                    md_parts.append(text)
            md_content = "\n\n".join(md_parts)
        else:
            md_content = self._strip_tags(html)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        self.report_progress(100, "Done")
        return True

    def _extract_text(self, path):
        html = self._read_html(path)
        if BS4_AVAILABLE:
            soup = BeautifulSoup(html, 'html.parser')
            for tag in soup(['script', 'style']):
                tag.decompose()
            return soup.get_text(separator='\n\n', strip=True)
        return self._strip_tags(html)

    def _strip_tags(self, html):
        text = re.sub(r'<script[^>]*>[\s\S]*?</script>', '', html, flags=re.IGNORECASE)
        text = re.sub(r'<style[^>]*>[\s\S]*?</style>', '', text, flags=re.IGNORECASE)
        text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
        text = re.sub(r'</(p|div|h[1-6]|li|tr)>', '\n\n', text, flags=re.IGNORECASE)
        text = re.sub(r'<[^>]+>', '', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()
