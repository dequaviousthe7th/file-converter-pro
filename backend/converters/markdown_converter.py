"""
Markdown Converter Module
Handles conversions from Markdown to other formats
"""

import os
import re
from pathlib import Path
from backend.converters.base_converter import BaseConverter, ConversionError

try:
    import pypandoc
    PYPANDOC_AVAILABLE = True
except ImportError:
    PYPANDOC_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except Exception:
    WEASYPRINT_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


class MarkdownConverter(BaseConverter):
    """Convert Markdown files to various formats"""

    SUPPORTED_OUTPUTS = ['pdf', 'docx', 'txt', 'html']

    def convert(self, input_path: str, output_path: str, target_format: str) -> bool:
        self.check_cancel()
        if target_format == 'pdf':
            return self._to_pdf(input_path, output_path)
        elif target_format == 'docx':
            return self._to_docx(input_path, output_path)
        elif target_format == 'html':
            return self._to_html(input_path, output_path)
        elif target_format == 'txt':
            return self._to_text(input_path, output_path)
        else:
            raise ConversionError(f"Unsupported target format: {target_format}")

    def _to_pdf(self, input_path: str, output_path: str) -> bool:
        self.report_progress(10, "Converting Markdown to PDF...")
        if WEASYPRINT_AVAILABLE:
            try:
                return self._to_pdf_weasyprint(input_path, output_path)
            except Exception:
                pass

        if REPORTLAB_AVAILABLE:
            return self._to_pdf_reportlab(input_path, output_path)

        if PYPANDOC_AVAILABLE:
            return self._convert_with_pandoc(input_path, output_path, 'pdf')

        raise ConversionError("No PDF library available",
                              "Install weasyprint or reportlab")

    def _to_pdf_reportlab(self, input_path: str, output_path: str) -> bool:
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                md_content = f.read()

            self.check_cancel()
            text = md_content
            text = re.sub(r'#{1,6}\s+', '', text)
            text = re.sub(r'\*\*|\*|__|_', '', text)
            text = re.sub(r'`', '', text)
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            text = re.sub(r'!\[([^\]]*)\]\([^\)]+\)', r'[Image: \1]', text)

            doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            title = Paragraph("Converted Document", styles['Title'])
            story.append(title)
            story.append(Spacer(1, 12))

            self.report_progress(50, "Building PDF...")
            for paragraph in text.split('\n\n'):
                self.check_cancel()
                if paragraph.strip():
                    safe = paragraph.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    p = Paragraph(safe.replace('\n', '<br/>'), styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 12))

            doc.build(story)
            self.report_progress(100, "Done")
            return True
        except ConversionError:
            raise
        except Exception as e:
            raise ConversionError(f"Markdown to PDF failed: {e}")

    def _to_pdf_weasyprint(self, input_path: str, output_path: str) -> bool:
        with open(input_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        self.check_cancel()
        self.report_progress(30, "Converting to HTML...")
        html_content = markdown.markdown(
            md_content, extensions=['tables', 'fenced_code', 'toc', 'nl2br']
        )

        css_style = """<style>
            @page { margin: 2cm; size: A4; }
            body { font-family: Arial, sans-serif; font-size: 11pt; line-height: 1.6; }
            h1 { font-size: 24pt; color: #1a1a1a; }
            h2 { font-size: 18pt; color: #2a2a2a; }
            code { background: #f4f4f4; padding: 2px 6px; }
            pre { background: #f4f4f4; padding: 16px; }
        </style>"""

        full_html = f'<!DOCTYPE html><html><head><meta charset="utf-8">{css_style}</head><body>{html_content}</body></html>'

        self.report_progress(60, "Generating PDF...")
        self.check_cancel()
        HTML(string=full_html).write_pdf(output_path)
        self.report_progress(100, "Done")
        return True

    def _to_docx(self, input_path: str, output_path: str) -> bool:
        self.report_progress(10, "Converting Markdown to Word...")
        if PYPANDOC_AVAILABLE:
            try:
                return self._convert_with_pandoc(input_path, output_path, 'docx')
            except Exception:
                pass

        if not PYTHON_DOCX_AVAILABLE:
            raise ConversionError("python-docx not installed")

        with open(input_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        self.check_cancel()
        doc = Document()

        for line in md_content.split('\n'):
            self.check_cancel()
            line = line.strip()
            if not line:
                continue
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            else:
                doc.add_paragraph(line)

        doc.save(output_path)
        self.report_progress(100, "Done")
        return True

    def _to_html(self, input_path: str, output_path: str) -> bool:
        if not MARKDOWN_AVAILABLE:
            raise ConversionError("markdown library not installed")

        self.report_progress(20, "Converting Markdown to HTML...")
        with open(input_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        self.check_cancel()
        html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])

        full_html = f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8"><title>Converted</title>
<style>body{{font-family:Arial,sans-serif;max-width:800px;margin:40px auto;line-height:1.6;}}</style>
</head><body>{html_content}</body></html>"""

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(full_html)
        self.report_progress(100, "Done")
        return True

    def _to_text(self, input_path: str, output_path: str) -> bool:
        self.report_progress(20, "Stripping Markdown formatting...")
        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()

        self.check_cancel()
        content = re.sub(r'```[\s\S]*?```', '', content)
        content = re.sub(r'`([^`]+)`', r'\1', content)
        content = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', content)
        content = re.sub(r'^#{1,6}\s+', '', content, flags=re.MULTILINE)
        content = re.sub(r'\*\*|__|\*|_', '', content)
        content = re.sub(r'^>\s?', '', content, flags=re.MULTILINE)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        self.report_progress(100, "Done")
        return True

    def _convert_with_pandoc(self, input_path: str, output_path: str, target_format: str) -> bool:
        self.check_cancel()
        pypandoc.convert_file(input_path, target_format, outputfile=output_path)
        self.report_progress(100, "Done")
        return True
