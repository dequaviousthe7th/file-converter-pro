"""
Text Converter Module
Handles conversions from plain text to other formats
"""

import os
from pathlib import Path
from backend.converters.base_converter import BaseConverter, ConversionError

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


class TextConverter(BaseConverter):
    """Convert text files to various formats"""

    SUPPORTED_OUTPUTS = ['pdf', 'docx', 'md']

    def convert(self, input_path: str, output_path: str, target_format: str) -> bool:
        self.check_cancel()
        if target_format == 'pdf':
            return self._to_pdf(input_path, output_path)
        elif target_format == 'docx':
            return self._to_docx(input_path, output_path)
        elif target_format == 'md':
            return self._to_markdown(input_path, output_path)
        else:
            raise ConversionError(f"Unsupported target format: {target_format}")

    def _to_pdf(self, input_path: str, output_path: str) -> bool:
        if not REPORTLAB_AVAILABLE:
            raise ConversionError("reportlab not installed")
        try:
            self.report_progress(10, "Reading text file...")
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

            self.check_cancel()
            self.report_progress(30, "Creating PDF...")

            doc = SimpleDocTemplate(output_path, pagesize=A4,
                                    rightMargin=72, leftMargin=72,
                                    topMargin=72, bottomMargin=18)
            styles = getSampleStyleSheet()
            normal = styles['Normal']
            normal.fontSize = 11
            normal.leading = 14
            story = []

            paragraphs = content.split('\n\n')
            total = len(paragraphs)
            for i, para_text in enumerate(paragraphs):
                self.check_cancel()
                if not para_text.strip():
                    story.append(Spacer(1, 12))
                    continue
                safe = para_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                safe = safe.replace('\n', '<br/>')
                story.append(Paragraph(safe, normal))
                story.append(Spacer(1, 12))
                self.report_progress(30 + int(60 * (i + 1) / max(total, 1)))

            doc.build(story)
            self.report_progress(100, "Done")
            return True
        except ConversionError:
            raise
        except Exception as e:
            raise ConversionError(f"Text to PDF failed: {e}")

    def _to_docx(self, input_path: str, output_path: str) -> bool:
        if not PYTHON_DOCX_AVAILABLE:
            raise ConversionError("python-docx not installed")
        try:
            self.report_progress(10, "Reading text file...")
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

            self.check_cancel()
            self.report_progress(30, "Creating Word document...")

            doc = Document()
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)

            for para_text in content.split('\n\n'):
                self.check_cancel()
                if not para_text.strip():
                    continue
                lines = para_text.split('\n')
                for i, line in enumerate(lines):
                    p = doc.add_paragraph(line)
                    if i < len(lines) - 1:
                        p.add_run().add_break()

            doc.save(output_path)
            self.report_progress(100, "Done")
            return True
        except ConversionError:
            raise
        except Exception as e:
            raise ConversionError(f"Text to DOCX failed: {e}")

    def _to_markdown(self, input_path: str, output_path: str) -> bool:
        try:
            self.report_progress(20, "Converting to Markdown...")
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

            self.check_cancel()
            filename = Path(input_path).stem
            md_content = f"# {filename}\n\n"

            for para in content.split('\n\n'):
                if para.strip():
                    para = para.replace('*', '\\*').replace('_', '\\_').replace('`', '\\`')
                    md_content += para + "\n\n"

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(md_content)
            self.report_progress(100, "Done")
            return True
        except ConversionError:
            raise
        except Exception as e:
            raise ConversionError(f"Text to Markdown failed: {e}")
